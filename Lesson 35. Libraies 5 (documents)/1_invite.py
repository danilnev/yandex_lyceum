from docx import Document
import sys

doc = Document()
place = input().strip().lower()
time = input().strip().lower()
peoples = list(map(str.rstrip, sys.stdin.readlines()))
for people in peoples:
    doc.add_paragraph(f'Приветствую, {people}!')
    doc.add_paragraph(f'Приглашаем вас на празднование 8 марта, {place}. Оно состоится {time}.')
    if people != peoples[-1]:
        doc.add_section()
doc.save('res.docx')
